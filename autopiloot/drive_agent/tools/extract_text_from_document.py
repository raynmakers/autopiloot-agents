"""
Extract clean text from various document formats
Robust text extraction pipeline for PDF, DOCX, plain text, and Google Workspace files
"""

import os
import json
import sys
import io
import re
from typing import Dict, Any, Optional, List
from pydantic import Field
from agency_swarm.tools import BaseTool
import base64

# Add config directory to path
from loader import get_config_value

# Optional imports for content extraction
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import html2text
    HTML2TEXT_AVAILABLE = True
except ImportError:
    HTML2TEXT_AVAILABLE = False


class ExtractTextFromDocument(BaseTool):
    """
    Extract clean, structured text from various document formats.
    Supports PDF, DOCX, plain text, HTML, CSV, and Google Workspace exports.
    """

    content: str = Field(
        ...,
        description="Document content (text content or base64 encoded binary data)"
    )

    mime_type: str = Field(
        ...,
        description="MIME type of the document (e.g., 'application/pdf', 'text/plain')"
    )

    file_name: str = Field(
        ...,
        description="Original file name for format detection and metadata"
    )

    content_encoding: str = Field(
        default="text",
        description="Content encoding: 'text' for plain text, 'base64' for binary data"
    )

    max_length: Optional[int] = Field(
        default=None,
        description="Maximum text length to extract (characters). If None, uses config default."
    )

    clean_text: bool = Field(
        default=True,
        description="Whether to clean and normalize extracted text"
    )

    extract_metadata: bool = Field(
        default=True,
        description="Whether to extract document metadata (page count, structure info)"
    )

    def _get_max_length(self) -> int:
        """Get maximum text length from config or parameter."""
        if self.max_length is not None:
            return self.max_length

        # Get from config - assuming reasonable default for RAG indexing
        drive_config = get_config_value("drive", {})
        tracking_config = drive_config.get("tracking", {})

        # Default to 50,000 characters (~25 pages of text)
        return tracking_config.get("max_text_length", 50000)

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not self.clean_text:
            return text

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove excessive newlines but preserve paragraph structure
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        # Remove common document artifacts
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        # Remove excessive punctuation
        text = re.sub(r'[.]{3,}', '...', text)
        text = re.sub(r'[-]{3,}', '---', text)

        # Clean up spacing around punctuation
        text = re.sub(r'\s+([,.!?;:])', r'\1', text)
        text = re.sub(r'([,.!?;:])\s+', r'\1 ', text)

        # Trim and normalize
        text = text.strip()

        return text

    def _extract_text_from_pdf(self, pdf_bytes: bytes) -> Dict[str, Any]:
        """Extract text and metadata from PDF bytes."""
        if not PDF_AVAILABLE:
            return {
                "text": "[PDF text extraction not available - PyPDF2 not installed]",
                "metadata": {"error": "PyPDF2 not available"}
            }

        try:
            pdf_file = io.BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = []
            page_texts = []

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        page_texts.append(page_text)
                        text_content.append(page_text)
                except Exception as e:
                    page_texts.append(f"[Error extracting page {page_num + 1}: {str(e)}]")

            # Combine all text
            full_text = "\n\n".join(text_content) if text_content else "[No text content found in PDF]"

            # Extract metadata
            metadata = {
                "page_count": len(pdf_reader.pages),
                "pages_with_text": len(page_texts),
                "extraction_method": "PyPDF2"
            }

            # Try to extract document info
            try:
                doc_info = pdf_reader.metadata
                if doc_info:
                    metadata.update({
                        "title": doc_info.get("/Title", "").strip() if doc_info.get("/Title") else None,
                        "author": doc_info.get("/Author", "").strip() if doc_info.get("/Author") else None,
                        "subject": doc_info.get("/Subject", "").strip() if doc_info.get("/Subject") else None,
                        "creator": doc_info.get("/Creator", "").strip() if doc_info.get("/Creator") else None
                    })
            except Exception:
                pass

            return {
                "text": full_text,
                "metadata": metadata
            }

        except Exception as e:
            return {
                "text": f"[Error processing PDF: {str(e)}]",
                "metadata": {"error": str(e), "extraction_method": "PyPDF2"}
            }

    def _extract_text_from_docx(self, docx_bytes: bytes) -> Dict[str, Any]:
        """Extract text and metadata from DOCX bytes."""
        if not DOCX_AVAILABLE:
            return {
                "text": "[DOCX text extraction not available - python-docx not installed]",
                "metadata": {"error": "python-docx not available"}
            }

        try:
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)

            text_content = []
            paragraph_count = 0

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    paragraph_count += 1

            # Combine all text
            full_text = "\n\n".join(text_content) if text_content else "[No text content found in DOCX]"

            # Extract metadata
            metadata = {
                "paragraph_count": paragraph_count,
                "extraction_method": "python-docx"
            }

            # Try to extract document properties
            try:
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title if core_props.title else None,
                    "author": core_props.author if core_props.author else None,
                    "subject": core_props.subject if core_props.subject else None,
                    "created": core_props.created.isoformat() if core_props.created else None,
                    "modified": core_props.modified.isoformat() if core_props.modified else None
                })
            except Exception:
                pass

            return {
                "text": full_text,
                "metadata": metadata
            }

        except Exception as e:
            return {
                "text": f"[Error processing DOCX: {str(e)}]",
                "metadata": {"error": str(e), "extraction_method": "python-docx"}
            }

    def _extract_text_from_html(self, html_content: str) -> Dict[str, Any]:
        """Extract text from HTML content."""
        if HTML2TEXT_AVAILABLE:
            try:
                h = html2text.HTML2Text()
                h.ignore_links = True
                h.ignore_images = True
                h.ignore_emphasis = False
                h.body_width = 0  # Don't wrap lines

                text = h.handle(html_content)

                return {
                    "text": text,
                    "metadata": {
                        "extraction_method": "html2text",
                        "original_length": len(html_content)
                    }
                }
            except Exception as e:
                pass

        # Fallback: simple HTML tag removal
        try:
            import re
            # Remove script and style elements
            text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)

            # Remove HTML tags
            text = re.sub(r'<[^>]+>', '', text)

            # Decode HTML entities
            import html
            text = html.unescape(text)

            return {
                "text": text,
                "metadata": {
                    "extraction_method": "regex_fallback",
                    "original_length": len(html_content)
                }
            }
        except Exception as e:
            return {
                "text": f"[Error processing HTML: {str(e)}]",
                "metadata": {"error": str(e)}
            }

    def _extract_text_from_csv(self, csv_content: str) -> Dict[str, Any]:
        """Extract structured text from CSV content."""
        try:
            import csv
            import io

            # Parse CSV
            csv_file = io.StringIO(csv_content)
            reader = csv.reader(csv_file)

            rows = list(reader)
            if not rows:
                return {
                    "text": "[Empty CSV file]",
                    "metadata": {"row_count": 0, "extraction_method": "csv"}
                }

            # Format as structured text
            text_lines = []

            # Add header if it exists
            if rows:
                header = rows[0]
                text_lines.append("Headers: " + ", ".join(header))
                text_lines.append("=" * 50)

                # Add data rows
                for i, row in enumerate(rows[1:], 1):
                    if i <= 100:  # Limit to first 100 rows for performance
                        row_text = " | ".join(str(cell) for cell in row)
                        text_lines.append(f"Row {i}: {row_text}")
                    else:
                        text_lines.append(f"... and {len(rows) - 101} more rows")
                        break

            full_text = "\n".join(text_lines)

            return {
                "text": full_text,
                "metadata": {
                    "row_count": len(rows),
                    "column_count": len(rows[0]) if rows else 0,
                    "extraction_method": "csv",
                    "truncated": len(rows) > 101
                }
            }

        except Exception as e:
            return {
                "text": f"[Error processing CSV: {str(e)}]",
                "metadata": {"error": str(e)}
            }

    def _get_content_bytes(self) -> bytes:
        """Get content as bytes, handling different encodings."""
        if self.content_encoding == "base64":
            return base64.b64decode(self.content)
        else:
            # Assume text content, encode to bytes
            return self.content.encode('utf-8')

    def run(self) -> str:
        """
        Extract text from document content.

        Returns:
            JSON string containing extracted text and metadata
        """
        try:
            extracted_data = {"text": "", "metadata": {}}

            # Handle different MIME types
            if self.mime_type.startswith('text/'):
                # Plain text files
                if self.content_encoding == "base64":
                    try:
                        content_bytes = self._get_content_bytes()
                        # Try different encodings
                        for encoding in ['utf-8', 'utf-16', 'latin-1']:
                            try:
                                text = content_bytes.decode(encoding)
                                break
                            except UnicodeDecodeError:
                                continue
                        else:
                            text = content_bytes.decode('utf-8', errors='replace')
                    except Exception as e:
                        text = f"[Error decoding text: {str(e)}]"
                else:
                    text = self.content

                # Handle CSV specifically
                if self.mime_type == 'text/csv' or self.file_name.lower().endswith('.csv'):
                    extracted_data = self._extract_text_from_csv(text)
                else:
                    extracted_data = {
                        "text": text,
                        "metadata": {
                            "extraction_method": "direct_text",
                            "mime_type": self.mime_type
                        }
                    }

            elif self.mime_type == 'application/pdf':
                # PDF extraction
                content_bytes = self._get_content_bytes()
                extracted_data = self._extract_text_from_pdf(content_bytes)

            elif self.mime_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            ]:
                # DOCX extraction
                content_bytes = self._get_content_bytes()
                extracted_data = self._extract_text_from_docx(content_bytes)

            elif self.mime_type in ['text/html', 'application/xhtml+xml']:
                # HTML extraction
                if self.content_encoding == "base64":
                    content_bytes = self._get_content_bytes()
                    html_content = content_bytes.decode('utf-8', errors='replace')
                else:
                    html_content = self.content

                extracted_data = self._extract_text_from_html(html_content)

            elif self.mime_type.startswith('application/vnd.google-apps'):
                # Google Workspace files (should be exported to text already)
                if self.content_encoding == "base64":
                    content_bytes = self._get_content_bytes()
                    text = content_bytes.decode('utf-8', errors='replace')
                else:
                    text = self.content

                extracted_data = {
                    "text": text,
                    "metadata": {
                        "extraction_method": "google_workspace_export",
                        "original_mime_type": self.mime_type
                    }
                }

            else:
                # Unsupported MIME type
                extracted_data = {
                    "text": f"[Text extraction not supported for MIME type: {self.mime_type}]",
                    "metadata": {
                        "error": f"Unsupported MIME type: {self.mime_type}",
                        "extraction_method": "unsupported"
                    }
                }

            # Clean the extracted text
            if extracted_data["text"] and not extracted_data["text"].startswith("["):
                extracted_data["text"] = self._clean_text(extracted_data["text"])

            # Apply length limit
            max_length = self._get_max_length()
            if len(extracted_data["text"]) > max_length:
                extracted_data["text"] = extracted_data["text"][:max_length] + "\n\n[Text truncated due to length limit]"
                extracted_data["metadata"]["truncated"] = True
                extracted_data["metadata"]["truncated_at"] = max_length

            # Add extraction metadata
            result = {
                "extracted_text": extracted_data["text"],
                "text_length": len(extracted_data["text"]),
                "file_info": {
                    "name": self.file_name,
                    "mime_type": self.mime_type,
                    "content_encoding": self.content_encoding
                }
            }

            if self.extract_metadata:
                result["document_metadata"] = extracted_data["metadata"]

            # Add text statistics
            if extracted_data["text"] and not extracted_data["text"].startswith("["):
                text = extracted_data["text"]
                result["text_stats"] = {
                    "character_count": len(text),
                    "word_count": len(text.split()),
                    "line_count": text.count('\n') + 1,
                    "paragraph_count": len([p for p in text.split('\n\n') if p.strip()])
                }

            return json.dumps(result)

        except Exception as e:
            return json.dumps({
                "error": "extraction_error",
                "message": f"Failed to extract text from document: {str(e)}",
                "details": {
                    "file_name": self.file_name,
                    "mime_type": self.mime_type,
                    "content_encoding": self.content_encoding,
                    "type": type(e).__name__
                }
            })


if __name__ == "__main__":
    # Test the tool
    print("Testing ExtractTextFromDocument tool...")

    # Test 1: Plain text content
    print("\n1. Testing plain text extraction...")
    tool = ExtractTextFromDocument(
        content="This is a test document.\n\nIt has multiple paragraphs.\n\nAnd some formatting.",
        mime_type="text/plain",
        file_name="test.txt",
        content_encoding="text"
    )
    result = tool.run()
    result_json = json.loads(result)
    print(f"Extracted text: {result_json.get('extracted_text')}")
    print(f"Text stats: {result_json.get('text_stats')}")

    # Test 2: CSV content
    print("\n2. Testing CSV extraction...")
    csv_content = "Name,Age,City\nJohn,25,NYC\nJane,30,LA\nBob,35,Chicago"
    tool = ExtractTextFromDocument(
        content=csv_content,
        mime_type="text/csv",
        file_name="test.csv",
        content_encoding="text"
    )
    result = tool.run()
    result_json = json.loads(result)
    print(f"Extracted text length: {result_json.get('text_length')}")
    print(f"Text preview: {result_json.get('extracted_text')[:200]}...")

    # Test 3: HTML content
    print("\n3. Testing HTML extraction...")
    html_content = "<html><body><h1>Title</h1><p>This is a paragraph.</p><p>Another paragraph.</p></body></html>"
    tool = ExtractTextFromDocument(
        content=html_content,
        mime_type="text/html",
        file_name="test.html",
        content_encoding="text"
    )
    result = tool.run()
    result_json = json.loads(result)
    print(f"Extracted text: {result_json.get('extracted_text')}")

    # Test 4: Unsupported format
    print("\n4. Testing unsupported format...")
    tool = ExtractTextFromDocument(
        content="binary content",
        mime_type="application/octet-stream",
        file_name="test.bin",
        content_encoding="text"
    )
    result = tool.run()
    result_json = json.loads(result)
    print(f"Result: {result_json.get('extracted_text')}")

    # Test 5: Base64 encoded content
    print("\n5. Testing base64 encoded text...")
    import base64
    original_text = "This is base64 encoded text content."
    encoded_content = base64.b64encode(original_text.encode('utf-8')).decode('ascii')

    tool = ExtractTextFromDocument(
        content=encoded_content,
        mime_type="text/plain",
        file_name="test_encoded.txt",
        content_encoding="base64"
    )
    result = tool.run()
    result_json = json.loads(result)
    print(f"Decoded text: {result_json.get('extracted_text')}")
    print(f"Metadata: {result_json.get('document_metadata')}")

    print("\nExtraction testing completed!")